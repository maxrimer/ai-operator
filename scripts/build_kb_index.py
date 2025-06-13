#!/usr/bin/env python3
# build_kb_index.py
"""
1. Читает все .docx из data/kb_docs/
2. Делит документ на чанки (абзацы + строки таблиц) через docx_to_chunks
3. Кидает в SentenceTransformer (multilingual-E5-large)
4. Сохраняет FAISS-Index + метаданные для RAG-поиска
"""

# ───── std lib ────────────────────────────────────────────────────────────────
import json, re, pathlib, argparse
from typing import List
# ───── third-party ────────────────────────────────────────────────────────────
import faiss, numpy as np, pandas as pd
from tqdm import tqdm
from docx import Document
from sentence_transformers import SentenceTransformer

# ───── config ─────────────────────────────────────────────────────────────────
DOC_DIR   = pathlib.Path("../../../Downloads/AI-суфлер общий доступ/все_кб_данные/")
OUT_DIR   = pathlib.Path("../data/processed")
OUT_META  = pathlib.Path("../data/processed/kb_meta.jsonl")
OUT_INDEX = pathlib.Path("../data/processed/kb_index.index")
MODEL     = "intfloat/multilingual-e5-large"
CHUNK_LEN = 120

# ───── helpers ────────────────────────────────────────────────────────────────
TOKEN_RE = re.compile(r"\s+")


def clean(txt: str) -> str:
    return TOKEN_RE.sub(" ", txt).strip()


def docx_to_chunks(path: pathlib.Path) -> List[str]:
    """Абзацы + строки таблиц, склеенные ' | '."""
    doc = Document(path)
    chunks = []

    # paragraphs
    for p in doc.paragraphs:
        txt = clean(p.text)
        if txt: chunks.append(txt)

    # tables
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            cells = [clean(c.text) for c in row.cells if clean(c.text)]
            if cells:
                chunks.append(" | ".join(cells))
    return chunks


def best_device() -> str:
    import torch
    if torch.cuda.is_available(): return "cuda"
    if torch.backends.mps.is_available(): return "mps"
    return "cpu"


# ───── main build ─────────────────────────────────────────────────────────────
def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    model = SentenceTransformer(MODEL, device=best_device())

    meta, embeds = [], []
    for doc_path in tqdm(DOC_DIR.glob("*.docx"), desc="⤵ parsing"):
        for idx, chunk in enumerate(docx_to_chunks(doc_path)):
            words = chunk.split()
            # скользящее окно по CHUNK_LEN
            for i in range(0, len(words), CHUNK_LEN):
                slice_ = " ".join(words[i:i+CHUNK_LEN])
                if len(slice_.split()) < 20:
                    continue
                emb = model.encode("passage: " + slice_,
                                   normalize_embeddings=True)
                embeds.append(emb)
                meta.append({
                    "doc":    doc_path.stem,
                    "chunk":  slice_,
                    "source": f"{doc_path.name}#{idx}:{i//CHUNK_LEN}"
                })

    vecs = np.asarray(embeds, dtype="float32")
    index = faiss.IndexFlatIP(vecs.shape[1])
    index.add(vecs)

    # — save —
    faiss.write_index(index, OUT_INDEX)
    with OUT_META.open("w", encoding="utf-8") as f:
        for row in meta:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")

    print(f"✔ saved {len(meta)} chunks → {OUT_INDEX}\n"
          f"🛈 мета: {OUT_META}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--rebuild", action="store_true",
                        help="пересобрать даже если файлы есть")
    args = parser.parse_args()

    if OUT_INDEX.exists() and OUT_META.exists() and not args.rebuild:
        print("Index already exists. Use --rebuild to overwrite.")
    else:
        build()
